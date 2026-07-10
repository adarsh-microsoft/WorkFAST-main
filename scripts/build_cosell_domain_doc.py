"""
Compose a comprehensive Word document containing the full business domain
content extracted from the 4 CoSell PowerPoint decks:
  1. GPSCRM Co-sell.pptx                 (CRM business fundamentals)
  2. CoSell_Domain_Business_Logic_Guide  (domain knowledge & business logic)
  3. CRM Cosell Deep dive.pptx            (data model, views, queries, security)
  4. CoSell 301.pptx                      (Partner Sharing report)

Nothing is omitted; content is reorganized into a logical reading flow with
proper Word styles, tables, glossaries, and code blocks.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DOCX = r"C:\Users\v-adevashish\OneDrive - Microsoft\Desktop\SupportAgentDomain\CoSell_Business_Domain.docx"
OUT_DOCX_COPY = r"c:\WorkFAST-main\generated-content\support-agent-domain\CoSell_Business_Domain.docx"

# ---- Palette -------------------------------------------------------------
MS_BLUE = RGBColor(0x00, 0x5A, 0x9E)
MS_DARK = RGBColor(0x1F, 0x1F, 0x1F)
ACCENT = RGBColor(0x0F, 0x6C, 0xBD)
GREY = RGBColor(0x59, 0x59, 0x59)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# ---- Base style ----------------------------------------------------------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = MS_DARK

# Tighten default spacing
pf = normal.paragraph_format
pf.space_after = Pt(6)
pf.line_spacing = 1.08


def _set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _shade_paragraph(p, hex_color):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def heading(text, level=1, color=MS_BLUE, page_break=False, space_before=12):
    if page_break:
        doc.add_page_break()
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.color.rgb = color
    run.font.name = "Calibri"
    if level == 1:
        run.font.size = Pt(20)
    elif level == 2:
        run.font.size = Pt(15)
    elif level == 3:
        run.font.size = Pt(12.5)
    else:
        run.font.size = Pt(11.5)
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after = Pt(4)
    h.paragraph_format.keep_with_next = True
    return h


def para(text="", bold=False, italic=False, size=11, color=MS_DARK,
         align=None, space_after=6, space_before=0):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.color.rgb = color
    return p


def lead(text):
    """Italic descriptive subheading line under a section title."""
    return para(text, italic=True, color=GREY, size=10.5, space_after=8)


def bullet(text, level=0, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(2)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.size = Pt(11)
        rest = p.add_run(text)
        rest.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
    return p


def numbered(text, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
    p.add_run(text)
    return p


def term(name, definition):
    """Glossary-style term + definition block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(name)
    r.bold = True
    r.font.color.rgb = ACCENT
    r.font.size = Pt(11)
    d = doc.add_paragraph()
    d.paragraph_format.left_indent = Inches(0.25)
    d.paragraph_format.space_after = Pt(6)
    dr = d.add_run(definition)
    dr.font.size = Pt(10.5)


def code_block(text, label=None):
    if label:
        lp = doc.add_paragraph()
        lp.paragraph_format.space_after = Pt(0)
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(10)
        lr.font.color.rgb = GREY
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent = Inches(0.1)
    _shade_paragraph(p, "F2F2F2")
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    # ensure east-asian font mapping too
    rpr = r._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Consolas")
    rfonts.set(qn("w:hAnsi"), "Consolas")
    return p


def table(headers, rows, col_widths=None, first_col_bold=False, caption=None):
    if caption:
        cp = doc.add_paragraph()
        cp.paragraph_format.space_after = Pt(2)
        cr = cp.add_run(caption)
        cr.bold = True
        cr.font.size = Pt(10.5)
        cr.font.color.rgb = GREY
    ncols = len(headers) if headers else len(rows[0])
    t = doc.add_table(rows=0, cols=ncols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    # Header
    if headers:
        hdr = t.add_row().cells
        for i, htext in enumerate(headers):
            hdr[i].text = ""
            hp = hdr[i].paragraphs[0]
            hr = hp.add_run(htext)
            hr.bold = True
            hr.font.size = Pt(10)
            hr.font.color.rgb = WHITE
            _set_cell_bg(hdr[i], "005A9E")
    # Body
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            cp = cells[ci].paragraphs[0]
            cr = cp.add_run("" if val is None else str(val))
            cr.font.size = Pt(9.5)
            if first_col_bold and ci == 0:
                cr.bold = True
            if not headers and ri % 2 == 1:
                _set_cell_bg(cells[ci], "F2F6FB")
            elif headers and ri % 2 == 1:
                _set_cell_bg(cells[ci], "EAF1FA")
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BBBBBB")
    pbdr.append(bottom)
    pPr.append(pbdr)


def callout(text, fill="FFF4CE", title=None):
    """Highlighted note box."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.cell(0, 0)
    _set_cell_bg(cell, fill)
    cell.text = ""
    if title:
        tp = cell.paragraphs[0]
        tr = tp.add_run(title)
        tr.bold = True
        tr.font.size = Pt(10.5)
        tr.font.color.rgb = RGBColor(0x7A, 0x5C, 0x00)
        bp = cell.add_paragraph()
        br = bp.add_run(text)
        br.font.size = Pt(10.5)
    else:
        tp = cell.paragraphs[0]
        tr = tp.add_run(text)
        tr.font.size = Pt(10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# =========================================================================
# COVER PAGE
# =========================================================================
for _ in range(3):
    doc.add_paragraph()

p = para("CoSell", bold=True, size=44, color=MS_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para("Business Domain Knowledge Compendium", bold=True, size=22, color=ACCENT,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
para("Domain Concepts · Business Logic · Data Model · Reporting Metrics",
     italic=True, size=12.5, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

para("Global Partner Solutions  |  MCAPS Data Engineering  |  FY26",
     bold=True, size=12, color=MS_DARK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

# Source decks box
t = doc.add_table(rows=1, cols=1)
t.style = "Table Grid"
t.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = t.cell(0, 0)
_set_cell_bg(cell, "EAF1FA")
cell.text = ""
hp = cell.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
hr = hp.add_run("Consolidated from 4 source presentations")
hr.bold = True
hr.font.size = Pt(11)
hr.font.color.rgb = MS_BLUE
for nm in [
    "1.  GPSCRM Co-sell  —  CRM business fundamentals (Leads, MCEM, Co-sell motions & requirements)",
    "2.  CoSell Domain Business Logic Guide  —  Domain knowledge, metric logic & stakeholder guide",
    "3.  CRM Cosell Deep Dive  —  Data model, GPS Mart views, queries & security",
    "4.  GPS 301 Partner Sharing  —  Partner Sharing report, metrics & architecture",
]:
    bp = cell.add_paragraph()
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = bp.add_run(nm)
    br.font.size = Pt(10.5)

doc.add_paragraph()
para("This document consolidates the complete business domain captured across the four CoSell "
     "training decks into a single reference. It is organized for progressive learning — from "
     "foundational CRM concepts, through deep domain business logic, into the underlying data model, "
     "and finally the Partner Sharing reporting product.",
     italic=True, size=10.5, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

# =========================================================================
# TABLE OF CONTENTS
# =========================================================================
heading("Contents", level=1, page_break=True)
toc_items = [
    ("Part 1 — GPS CRM Co-sell: Business Fundamentals", 0),
    ("Leads and Opportunities", 1),
    ("Billed vs Consumption Pipelines", 1),
    ("Microsoft Customer Engagement Methodology (MCEM)", 1),
    ("What is Co-selling? — The Four Co-sell Types", 1),
    ("The Co-sell Motion (Inbound & Outbound)", 1),
    ("Co-sell Requirements", 1),
    ("IP Co-sell Incentive Program", 1),
    ("Azure IP Co-sell Process", 1),
    ("Marketplace Billed Sales (MBS) Co-sell", 1),
    ("Appendix: ARC & FY24 Metric Definitions", 1),
    ("Part 2 — CoSell Domain: Business Logic & Stakeholder Guide", 0),
    ("What is CoSell?", 1),
    ("Key Business Terminology", 1),
    ("The Co-Sell Deal Lifecycle", 1),
    ("Deal Directions: Inbound vs Outbound", 1),
    ("Currency Conversion Logic", 1),
    ("Deal Classification Rules", 1),
    ("Velocity & Performance Metrics", 1),
    ("Partner Identity Resolution", 1),
    ("Scorecard Recognition Logic", 1),
    ("Geography Resolution Chain", 1),
    ("IP CoSell, ISV Connect & Marketplace", 1),
    ("CoMarketing Module", 1),
    ("Data Architecture: Source to Gold", 1),
    ("Key Business Questions CoSell Answers", 1),
    ("Summary & Key Takeaways", 1),
    ("Part 3 — CRM Co-sell Deep Dive: Data & Technical", 0),
    ("Upstream Sources & GPS Mart", 1),
    ("GPS Mart Views and Tables", 1),
    ("GPS Mart ER Diagram (Star Schema)", 1),
    ("Important Queries", 1),
    ("CRM Tabular Model", 1),
    ("Marketplace Security Implementation", 1),
    ("Part 4 — GPS 301: Partner Sharing Report", 0),
    ("Partner Sharing Report Overview", 1),
    ("Billed Pipeline Metrics", 1),
    ("Consumption Pipeline Metrics", 1),
    ("FY26 Partner Sharing Metrics", 1),
    ("Reporting Structure, Slicers & Architecture", 1),
    ("Reporting Architecture Glossary", 1),
]
for label, lvl in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    if lvl == 0:
        p.paragraph_format.space_before = Pt(6)
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(11.5)
        r.font.color.rgb = MS_BLUE
    else:
        p.paragraph_format.left_indent = Inches(0.3)
        r = p.add_run("•  " + label)
        r.font.size = Pt(10.5)
        r.font.color.rgb = MS_DARK

# =========================================================================
# PART 1 — GPS CRM CO-SELL (Business Fundamentals)
# =========================================================================
heading("Part 1 — GPS CRM Co-sell: Business Fundamentals", level=1, page_break=True)
lead("Source deck: GPSCRM Co-sell — GPS Domain Training (CRM). Foundational concepts of leads, "
     "opportunities, the Microsoft engagement methodology, and the co-sell motions, requirements, "
     "and incentive programs.")

heading("Leads and Opportunities", level=2)
term("Lead", "Leads appear in CRM Online from several sources. Sellers may create them, or one can "
     "upload them from marketing campaigns via Worldwide Lead Management or other channels.")
term("Billed Opportunity", "An opportunity represents an engagement with a customer. Opportunities "
     "identify a customer's business needs, willingness to fund, purchasing timeline, and use of a "
     "primary sponsor.")
term("Billed Revenue", "All monetary transactions of Microsoft's customers are managed through "
     "Billed Opportunity.")
term("Consumption Opportunity", "Tracks the progress of projects that drive consumed revenue.")

heading("Billed vs Consumption Pipelines", level=2)
bullet("ACR (Azure Consumed Revenue) Pipeline is managed using Consumption Opportunities.")
bullet("Billed Revenue Pipeline is managed using Billed Opportunities.")

heading("Microsoft Customer Engagement Methodology (MCEM)", level=2)
lead("MCEM frames every customer engagement as a value-realization journey. Each opportunity "
     "(Billed or Consumption) advances through milestones, attaching Licensing, Support, and "
     "Consulting products along the way.")
para("MCEM defines five engagement phases, each with the Microsoft role, the customer intention / "
     "outcome, and the exit criteria that move the engagement forward. The process runs under a "
     "Continuous Account Plan.", space_after=8)
table(
    ["Phase", "Microsoft Role", "Intention / Customer Outcome", "Exit Criteria"],
    [
        ["Listen & Consult", "Help qualify and quantify the customer needs",
         "Customer believes Microsoft could solve the problem", "Discover Opportunity"],
        ["Inspire & Design", "Inspire & design the solution with the customer",
         "Customer has interest in the proposed solution", "Define business value and solution"],
        ["Empower & Achieve", "Prove business case and technology",
         "Customer has confidence in the solution", "Prove and gain decision on the solution"],
        ["Deliver (Connected Delivery)", "Support delivery of the solution aligned to business value",
         "Customer satisfied with implementation", "Deliver on solution and realize value"],
        ["Drive Continuous Innovation (Manage & Optimize)", "Monitor, optimize, and expand value",
         "Customer expectations are met", "Optimize & identify expansion opportunity"],
    ],
    col_widths=[1.5, 1.7, 1.9, 1.5],
)
callout("Value Realization Process is WHAT we will do; Value Realization Methodology is HOW we will "
        "do it. Both run continuously under the Continuous Account Plan. (Microsoft Confidential)",
        title="Key distinction")

heading("What is Co-selling? — The Four Co-sell Types", level=2)
para("Co-selling with Microsoft is defined as any collaborative engagement between Microsoft and our "
     "partner ecosystem — including building demand, sales planning, sharing sales leads, accelerating "
     "partner-to-partner empowered selling, and delivering marketplace-led commerce for customer "
     "transformation.", space_after=8)
term("Sell through partner", "Partners sell / resell Microsoft technology through their own sellers.")
term("Services Co-sell", "Sellers and partners collaborate to sell a partner's project / managed "
     "services with Microsoft technology embedded.")
term("IP Co-sell", "Sellers and partners collaborate to sell a partner's repeatable, packaged IP "
     "solution.")
term("Partners selling partners", "Partners sell / resell another partner's offer or solution to the "
     "customer.")

heading("The Co-sell Motion (Inbound & Outbound)", level=2)
lead("Co-sell deals flow between two portals: Partner Center (for Microsoft Partners) and MSX (for "
     "Microsoft Sellers). Direction depends on who initiates.")
para("Inbound (Partner → Microsoft):", bold=True, space_after=2)
bullet("Pipeline is created by the Partner first; a Deal is created in Partner Center.")
bullet("Partner Led: the partner works on a deal alone.")
bullet("Active Co-sell: the partner shares the deal with Microsoft.")
bullet("Pending: the seller has not yet responded to the deal.")
bullet("Expired: the seller does not respond within 14 days.")
bullet("If the seller accepts the deal, an Opportunity is created in MSX.")
para("Outbound (Microsoft → Partner):", bold=True, space_after=2, space_before=6)
bullet("The seller creates an Opportunity and selects a partner as draft.")
bullet("Microsoft shares with the partner; a Partner Deal is created in Partner Center.")

heading("Co-sell Requirements", level=2)
lead("Eligibility and channel enablement requirements by offer type.")
table(
    ["Offer Type", "Offer Requirements", "Offer Status", "Channel Enablement"],
    [
        ["All",
         "Meet the Microsoft commercial marketplace publishing requirements. Create a business "
         "profile in Partner Center.",
         "In Market",
         "Offer made available in Azure Marketplace and/or AppSource; business profile available in "
         "Solution Provider; direct customer purchasing and referrals."],
        ["IP",
         "Provide partner sales contact for each eligible geo. Upload required bill of materials: "
         "(1) customer one-pager, (2) customer presentation deck. Optional: (3) solution landing "
         "page URL, (4) social media channel URL or other (e.g., case study).",
         "Co-sell Ready",
         "Offer made available in Microsoft seller catalog; offer sellable by another partner; "
         "ability to share co-sell opportunities with Microsoft or partner sellers."],
        ["Services",
         "Provide partner sales contact for each eligible geo. Upload required bill of materials: "
         "(1) customer one-pager, (2) customer presentation deck. Optional: (3) solution landing "
         "page URL, (4) social media channel URL or case study. Solution Partner Designation in at "
         "least one Solution Area.",
         "Co-sell Ready",
         "Offer made available in Microsoft seller catalog; offer sellable by another partner; "
         "ability to share co-sell opportunities with Microsoft or partner sellers."],
    ],
    col_widths=[0.9, 2.6, 1.0, 2.1],
)
para("Incentive eligibility — technical validation:", bold=True, space_before=6, space_after=2)
bullet("Transactable offers (Get It Now, Free Trial — Azure VMs, Azure Application) are "
       "automatically made incentive eligible upon verification of offer type. Technical validation "
       "is built into publication; no further manual validation is required.")
bullet("Non-Transactable offers (List, Contact Me, Test Drive, SaaS) require manual Azure IP "
       "Technical Validation.")
table(
    ["Offer Type", "Offer Requirements", "Offer Status", "Channel Enablement"],
    [
        ["Azure IP — incentive eligible",
         "Meet or exceed $100k USD Azure Consumed Revenue (ACR) TTM (trailing twelve months) "
         "threshold at an organizational level, OR reach the $100k USD billed revenue threshold "
         "with a transactable offer. Pass the Microsoft technical validation for an Azure-based "
         "solution. Provide a reference architecture diagram.",
         "Co-sell Incentivized",
         "Microsoft seller receives incentive credit if offer is sold; preferred solution badging "
         "in Azure Marketplace and/or AppSource."],
        ["Business Applications — incentive eligible",
         "Active enrollment in the ISV Connect program.",
         "Co-sell Incentivized",
         "Microsoft seller receives incentive credit if offer is sold; preferred solution badging "
         "in Azure Marketplace and/or AppSource."],
    ],
    col_widths=[1.3, 2.8, 1.1, 2.0],
)

heading("IP Co-sell Incentive Program", level=2)
para("Intent of the program: encourage collaborative co-selling between Microsoft sellers and "
     "Microsoft IP solution partners — accelerating Azure usage, driving digital transformation, and "
     "fostering long-term partnership.", space_after=6)
para("Benefits:", bold=True, space_after=2)
bullet("Partners bring industry expertise and the finished solutions that today's customers are "
       "looking for.")
bullet("Sellers leverage partner-built solutions via the co-sell model to continue meeting customer "
       "success and Microsoft consumption goals.")
callout("\u201cLeading companies in every industry are partnering with us to build their own digital "
        "capability to compete and grow. This is creating broader opportunity for everyone, including "
        "our ecosystem.\u201d  — Satya Nadella", fill="EAF1FA")

heading("Azure IP Co-sell Process", level=2)
lead("The end-to-end co-sell sales motion, from lead sharing to incentive credit.")
numbered("", bold_lead="Lead Shared — ")
para("Partners or Microsoft sellers initiate the process by sharing opportunities.", space_after=4)
numbered("", bold_lead="Co-sell Sales Process — ")
para("Sales process occurs via Active Co-sell or Partner-led Co-sell with an Incentive-Eligible "
     "Solution (driven by a partner or field seller).", space_after=4)
numbered("", bold_lead="Deal Registered by Partner — ")
para("After customer contract signature, deals are set to 100% in Partner Center.", space_after=4)
numbered("", bold_lead="Validation & Activation — ")
para("Deal Validation + Deal Review. The deal goes through validation and review. Partners may need "
     "to share proof of execution as part of the deal validation process.", space_after=4)
numbered("", bold_lead="Scorecard + Incentive Credited — ")
para("Earn a percentage or flat rate per deal based on the ACV and where the solution is deployed.",
     space_after=4)

heading("Marketplace Billed Sales (MBS) Co-sell", level=2)
para("Metric definition: the amount Microsoft invoices customers or CSPs for the Azure IP Co-sell "
     "offers purchased through Microsoft Commercial Marketplace.", space_after=4)
bullet("Annualized and capped at $1M maximum.")
bullet("Excludes invoices in China and Korea, or by partners based in China.")
bullet("Excludes billing from US Fed and all EDU and Public Sector / Government segments.")
para("Two crediting lenses:", bold=True, space_before=6, space_after=2)
term("MBS Co-sell — Commercial", "Sellers receive quota and credit. Only applies to Top-Tier "
     "partners. Targets are set by Sub Region and Sub Segment.")
term("MBS Co-sell — GPS", "PDMs receive quota and credit. Only applies to partners included in the "
     "MPL. Targets are set by PartnerOne ID for the WW MBS driven by each respective partner.")
callout("IPCS Transition Program: 62 identified partners were granted an extension to transition to "
        "Marketplace. Permission was granted to continue the traditional IP Co-sell program through "
        "Deal Reg into FY24. This program is expected to retire in FY25. Sellers receive quota and "
        "credit based on a new metric called ARC (Artificial Revenue Credit) — see Appendix. "
        "(Microsoft Confidential)", title="Transition note")

heading("Appendix: ARC & FY24 Metric Definitions", level=2)
para("MBS Component — IP Co-Sell Partners Marketplace Transition (CrossAzure):", bold=True,
     space_after=2)
para("Artificial Revenue Credit (ARC) calculation:", space_after=2)
bullet("SaaS: ARC of 30% of ACV or $15K, whichever is higher (payout rates). ARC capped at $300K "
       "maximum per deal.")
bullet("BYOL: ARC of 5%, capped at $50K.")
para("(Microsoft Internal Only)", italic=True, size=9.5, color=GREY, space_after=8)

para("FY24 Metric Deck — Section 1: Definition and Scope (ARC / Deal Registration ACR):",
     bold=True, space_after=2)
table(
    ["Attribute", "Detail"],
    [
        ["One-line definition",
         "$ Deal Registration ACR. Limited co-sell program retires ACR comp bucket as artificial "
         "revenue credit, with related quota to all eligible sellers (mechanics similar to FY23)."],
        ["Detailed definition",
         "Artificial Revenue Credit (ARC) as calculated from Annual Contract Value (ACV) from won "
         "Enterprise & SMC-Corporate segment deals incorporating an Azure incentive-eligible "
         "repeatable IP solution, registered by a partner using Partner Center. Minimum deal size of "
         "$25,000 USD of contract value within the first year. ACV auto-calculated for multi-year "
         "contracts. Approved ARC determined by Sales Operations through deal validation: 30% of ACV "
         "or $15k (whichever is higher, capped at $300K) for bundled deals; 5% of ACV (capped at "
         "$50K) for BYOL deals."],
        ["Business objective",
         "Preserve and safeguard revenue by providing the highest-ACV Top-Tier partners a 6-month "
         "transition period wherein sellers continue to receive ACR retirement for co-selling "
         "outside of Marketplace."],
        ["In scope",
         "End customer must be Enterprise segment or SMC-Corporate. In-scope deals are those where "
         "co-sell occurs outside of Marketplace. Solutions must be listed on Marketplace (need not "
         "be transactable) and have the \u201cIPCS Incentivized\u201d designation (technically "
         "reviewed, drives material Azure consumption across partner and customer tenants)."],
        ["Out of scope",
         "Total Contract Value excludes Microsoft first-party licenses, non-related services, and "
         "any hardware. Renewal of subscriptions / licenses also excluded."],
        ["YoY changes",
         "New metric component. Scope limited to a select list of top-performing ACV partners — a "
         "subsegment of Top Tier."],
        ["Upstream data source", "Starlight / AHR."],
    ],
    col_widths=[1.5, 5.4], first_col_bold=True,
)
para("Segment Scope (ARC):", bold=True, space_before=6, space_after=2)
table(
    ["Enterprise Segment", "In Scope", "SMC Segment", "In Scope"],
    [
        ["Major Commercial", "Yes", "Enterprise Growth (China only)", "Yes"],
        ["Strategic Commercial", "Yes", "SM&C Commercial - Corporate", "Yes"],
        ["Major Public Sector", "Yes", "SM&C Education - Corporate", "Yes"],
        ["Strategic Public Sector", "Yes", "SM&C Government - Corporate", "Yes"],
        ["Education (under Major Public Sector)", "Yes", "SM&C Commercial - SMB", "No"],
        ["", "", "SM&C Education - SMB", "No"],
        ["", "", "SM&C Government - SMB", "No"],
    ],
    col_widths=[2.4, 0.9, 2.5, 0.9],
)
para("Compensation alignment (ARC): Metric Owner — Laurie Viault / Mason McCoy;  Finance Contact — "
     "Avi Ardon;  WWIC — Brenda Bruce.", size=10, color=GREY, space_before=4)

para("FY24 Metric Deck — Marketplace Billed Sales (MBS):", bold=True, space_before=8, space_after=2)
table(
    ["Attribute", "Detail"],
    [
        ["One-line definition",
         "The amount Microsoft invoices customers or CSPs for the Azure IP co-sell offers purchased "
         "through Microsoft Commercial Marketplace."],
        ["Detailed definition",
         "After customers or CSPs purchase Azure IP Co-Sell tagged solutions in Microsoft Commercial "
         "Marketplace, Microsoft invoices them. MBS appears only once Microsoft issues an invoice. "
         "Includes all Azure IP Co-sell tagged solutions from Top-Tier partners that are "
         "transactable and where Microsoft invoices the customer — Azure applications (Virtual "
         "Machine Images, Solution Templates, SaaS). Seller credit: MBS annualized value, 1x credit "
         "for Top Tier with $1M cap; cap per deal applied after modifier for annualized & monthly. "
         "Sellers receive quota and actuals."],
        ["Business objective",
         "Measures the sale of 3rd-party IP co-sell tagged solutions that run on Microsoft Cloud, "
         "growing Azure cloud consumption. The incentive is a key MCPP benefit for Top-Tier IP "
         "partners, who gain access to Managed Account Teams. Limiting to Top-Tier incents partners "
         "to meet criteria and join the Top-Tier list."],
        ["In scope",
         "All 3rd-party software applications that are Azure IP Co-sell incentivized from Top-Tier "
         "partners and transactable in Microsoft Commercial Marketplace. Includes direct, channel, "
         "and field MBS."],
        ["Out of scope",
         "All products outside the Azure IP Co-Sell solution tag. IP software only — no services "
         "offers. Non-SaaS Teams; any Contact Me; BYOL. Out-of-scope segment: Azure Gov. "
         "Out-of-scope geos: China, South Korea (Marketplace commerce only available for non-EA "
         "customers; IPCS transactions with EA customers not covered via MBS)."],
        ["YoY changes", "New metric."],
        ["Upstream actuals source", "MHR (planning), MECH (comp). Credited geo based on Partner Geo."],
        ["Partner association method", "Marketplace Billed Revenue (all else excluded)."],
        ["Business guidance",
         "Encourage the use of Marketplace as the standard platform for going to market with "
         "Microsoft, where partners publish offers and customers can purchase them. Metric Owner — "
         "Laurie Viault / Mason McCoy; Finance Contact — Avi Ardon; WWIC — Brenda Bruce."],
    ],
    col_widths=[1.5, 5.4], first_col_bold=True,
)

# =========================================================================
# PART 2 — COSELL DOMAIN BUSINESS LOGIC GUIDE
# =========================================================================
heading("Part 2 — CoSell Domain: Business Logic & Stakeholder Guide", level=1, page_break=True)
lead("Source deck: CoSell Data Stream — Domain Knowledge, Business Logic & Stakeholder Guide "
     "(Global Partner Solutions | MCAPS Data Engineering | FY26). This explains the business meaning "
     "behind CoSell — what the domain represents, how key decisions are encoded in data, and the "
     "logic that drives every metric.")

heading("What is CoSell?", level=2)
para("The Business Problem", bold=True, color=ACCENT, space_after=2)
para("Microsoft partners (ISVs, SIs, resellers) co-sell Microsoft products alongside their own "
     "solutions. Tracking whether a partner referral turned into a real opportunity, whether revenue "
     "was recognized, and how fast deals progressed through the funnel requires combining data from "
     "multiple disconnected systems (Partner Center, MSX/CRM, currency tables, account hierarchies).")
para("What CoSell Does", bold=True, color=ACCENT, space_after=2)
para("CoSell unifies this data into a single analytical model that answers:", space_after=2)
bullet("Did a partner referral convert into a Microsoft sales opportunity?")
bullet("How fast did Microsoft respond (inbound velocity)?")
bullet("What is the deal worth in USD across 9 different financial measures?")
bullet("Was the partner's CRM account co-sell-ready (ISV eligible)?")
bullet("Which scorecard month gets fiscal credit for this deal?")
bullet("What is the partner's quantitative impact on Microsoft's business (PIN)?")
para("Who Depends on CoSell", bold=True, color=ACCENT, space_after=2, space_before=4)
para("GPS Leadership, Partner Development Managers, Finance (revenue attribution), Field Sellers "
     "(scorecard credit), ISV Program (co-sell readiness tracking), and CoMarketing (investment ROI).")
para("CoSell at a glance:", bold=True, space_before=6, space_after=2)
table(
    ["Dimension", "Scale / Description"],
    [
        ["80+ Gold Tables", "Dimensional model covering the full deal lifecycle."],
        ["9 USD Financial Metrics", "With sophisticated multi-currency exchange rate logic."],
        ["12 Performance Metrics", "Velocity, acceptance, ISV engagement, win rates."],
        ["5 Source Systems Unified", "Partner Center, MSX/CRM, Hub, Currency, Revenue Account."],
        ["PIN Metrics", "Partner Impact Numbers across E5 Security, CSP Copilot, etc."],
    ],
    col_widths=[2.1, 4.8], first_col_bold=True,
)

heading("Key Business Terminology", level=2)
lead("The language of CoSell — terms every stakeholder must know.")
term("Co-Sell", "A joint selling motion where Microsoft and a partner collaborate to sell to a "
     "customer. The partner refers a deal (or Microsoft refers to the partner), and both work "
     "together through the sales cycle.")
term("Referral / Partner Deal", "A deal initiated in Partner Center by a partner (Outbound from the "
     "MS perspective) or sent to a partner by Microsoft (Inbound). Each gets a unique PSXDealID. This "
     "is the atomic unit of CoSell tracking.")
term("PartnerOne ID", "Microsoft's unified partner identity. A single partner may have multiple MPN "
     "IDs, but PartnerOne consolidates them into one hierarchy: PartnerOneID → PartnerOneSubKey "
     "(subsidiary level) → PartnerOneReportingKey (rolled up for reporting).")
term("Partner Sales Stage", "The lifecycle stage of a deal: Support 100%, Deploy 100%, MS Win 100% "
     "(positive closures), or Disengaged 0%, Loss 0% (negative closures). A deal is 'Closed' when it "
     "reaches any of these terminal stages.")
term("ISV Co-Sell Ready", "A flag on the partner's CRM account (IsCoSellReadyPartnerAccount). Only "
     "deals from ISV co-sell-ready partners are counted in ISVDealsReporting. This is an eligibility "
     "gate, not a deal attribute.")
term("Partner GTM (Go-To-Market)", "A flag (IsDealWithPartnerGTM) indicating the deal participates "
     "in Microsoft's Partner GTM program. Used as a prerequisite for IsDealAccepted and "
     "IsOpportunityWon calculations.")
term("PIN (Partner Impact Number)", "A quantitative measure of partner contribution to Microsoft's "
     "business. Metrics like E5 Security Usage and CSP Copilot Seats are per-partner (Non-Aggregate); "
     "others can be summed across partners (Aggregate).")
term("ACR (Azure Consumed Revenue)", "The actual Azure consumption revenue attributed to partner "
     "deals. PartnerReportedACR is self-reported by partners; TrueACRConsumption is system-verified.")
term("Scorecard Recognition Month", "The fiscal month a deal gets credited to. Governed by the "
     "3rd-of-month rule: deals approved after the 3rd → current month; before the 3rd → prior month.")
term("ATU (Account Team Unit)", "Microsoft's internal sales org structure. Mapped via MSSalesID → "
     "ATUID from DimRevenueAccount. Connects deals to the field sales team responsible.")

heading("The Co-Sell Deal Lifecycle", level=2)
lead("End-to-end journey from referral to revenue — and the data captured at each stage.")
table(
    ["Stage", "What Happens", "Data Captured"],
    [
        ["1. Referral Creation",
         "Partner creates a deal in Partner Center (or Microsoft refers to a partner).",
         "PSXDealID, Direction (Inbound/Outbound), CreatedDate, MPNId"],
        ["2. Acceptance & Linking",
         "Microsoft or Partner accepts. Deal linked to an MSX Opportunity.",
         "AcceptanceStatus, OpportunityID, CRMPartnerAccountKey"],
        ["3. Registration & Approval",
         "Deal registered and approved for scorecard. Oct-2018 cutoff applies for FX rates.",
         "RegistrationDate, ApprovalDate, Status"],
        ["4. Deal Progression",
         "Deal moves through sales stages. Velocity metrics computed.",
         "PartnerSalesStage, ApprovalVelocity, CloseVelocity"],
        ["5. Revenue Recognition",
         "Financial values converted to USD. Won/Closed status determined.",
         "DealValueCD, PartnerRevenueCD, BilledRevenueStatus"],
    ],
    col_widths=[1.5, 2.9, 2.5], first_col_bold=True,
)
para("Critical business rules applied during the lifecycle:", bold=True, space_before=6, space_after=2)
code_block(
    "ClosedPartnerDeals = 1  WHEN PartnerSalesStage IN\n"
    "  ('Support 100%', 'Deploy 100%', 'MS Win 100%', 'Disengaged 0%', 'Loss 0%')\n\n"
    "PartnerDealsAcceptance = 1  WHEN OpportunityID linked + valid PartnerDealID\n"
    "  + known CRM Partner + AcceptanceStatus = 'Accepted'\n\n"
    "ISVDealsReporting = 1  WHEN OpportunityID linked\n"
    "  + CRM partner's IsCoSellReadyPartnerAccount = 'true'\n\n"
    "IsDealAccepted = 'Yes'  WHEN IsDealWithPartnerGTM = 1\n"
    "  AND (Microsoft or Partner sub-status = 'Accepted')\n\n"
    "IsOpportunityWon = 'Yes'  WHEN IsDealWithPartnerGTM = 1\n"
    "  AND (BilledRevenueStatus = 'Won'\n"
    "       OR ConsumptionStatus = 'Closed' + Reason = 'Completed')"
)

heading("Deal Directions: Inbound vs Outbound", level=2)
lead("How co-sell deals flow between Microsoft and Partners.")
para("INBOUND  (Partner → Microsoft)", bold=True, color=ACCENT, space_after=2)
para("The partner identifies a customer opportunity and refers it TO Microsoft through Partner "
     "Center. Business meaning: the partner is bringing Microsoft into their deal. Velocity metric: "
     "VelocityOfInboundDeals measures how many days it takes Microsoft to create an MSX Opportunity "
     "after receiving the partner's referral.", space_after=2)
code_block(
    "IF Direction = 'Inbound' AND OpportunityCreatedDate IS NOT NULL:\n"
    "    VelocityOfInboundDeals = DATEDIFF(OpportunityCreatedDate, CreatedDate)\n"
    "    (minimum 1 day if same-day)\n"
    "ELSE: 0  (not applicable)"
)
para("OUTBOUND  (Microsoft → Partner)", bold=True, color=ACCENT, space_after=2)
para("Microsoft identifies an opportunity and refers it TO a partner through Partner Center. "
     "Business meaning: Microsoft is inviting a partner to collaborate on their deal. Velocity is not "
     "tracked via VelocityOfInboundDeals (= 0 for outbound); instead, PartnerDealsAcceptance tracks "
     "whether the partner accepted the referral.", space_after=2)
code_block(
    "PartnerDealsAcceptance = 1 WHEN:\n"
    "  - OpportunityID exists (deal was linked)\n"
    "  - PartnerDealID is valid\n"
    "  - CRM Partner is known (PartnerName IS NOT NULL)\n"
    "  - PartnerAcceptanceStatus = 'Accepted'"
)
callout("Private deals (ReferralSubType = 'private - pc') have special geography handling: "
        "CustomerGeographyKey is forced to 0 and ReportingGeographyKey to -9999. The "
        "PartnerOneReportingKey falls back to the Bridge table with SubsidiaryID = -9999.",
        title="Private deal exception")

heading("Currency Conversion Logic", level=2)
lead("The Oct-2018 cutoff rule and multi-source exchange rate fallback strategy.")
para("All financial metrics are stored in USD. Because deals originate in local currencies from "
     "Partner Center and MSX, the system uses exchange rates from multiple sources with a "
     "date-dependent priority order.")
callout("THE OCT-2018 RULE:  RegistrationDate < 2018-10-01 → Opportunity rate preferred first.  |  "
        "RegistrationDate ≥ 2018-10-01 → Partner Deal rate preferred first.", fill="EAF1FA",
        title="Date cutoff")
para("Simple conversion metrics (single source):", bold=True, space_before=4, space_after=2)
table(
    ["Metric", "Formula / Source"],
    [
        ["DealValueCD", "DealValue / CurrencyExchangeRate (from Currency table by CurrencyID)"],
        ["PartnerDealValueCD", "PartnerDealValue / rate by PartnerDealCurrencyID"],
        ["MicrosoftDealValueCD", "MicrosoftDealValue / rate by MicrosoftDealCurrencyID"],
        ["SolutionValueCD", "SolutionValue / rate by SolutionCurrencyID"],
    ],
    col_widths=[2.0, 4.9], first_col_bold=True,
)
para("Date-dependent fallback metrics:", bold=True, space_before=6, space_after=2)
table(
    ["Metric", "Fallback Chain"],
    [
        ["ProductRevenueCD (Dual fallback)",
         "Pre-Oct 2018: Opportunity rate → Partner Deal rate → 1.  "
         "Post-Oct 2018: Partner Deal rate → Opportunity rate → 1."],
        ["PartnerRevenueCD (Triple fallback)",
         "Pre-Oct 2018: Opportunity → Partner Deal → Currency table → 1.  "
         "Post-Oct 2018: Partner Deal → Currency table → Opportunity → 1."],
        ["ContractValueCD (Triple fallback)", "Same as PartnerRevenueCD."],
        ["EstimatedRevenueCD (Dual fallback)", "Same as ProductRevenueCD."],
    ],
    col_widths=[2.3, 4.6], first_col_bold=True,
)
callout("WHY THIS MATTERS: The fallback chain ensures no deal has NULL financial values. If the "
        "preferred exchange rate source is missing (e.g., Opportunity has no currency code), the "
        "system gracefully falls back to the next source. The Oct-2018 cutoff reflects a historical "
        "change in how Partner Center recorded currency vs. how MSX did — older deals trust "
        "Opportunity rates more; newer deals trust Partner Deal rates.")

heading("Deal Classification Rules", level=2)
lead("How deals get categorized — the exact business logic behind each flag.")
table(
    ["Flag", "Logic", "Business Meaning"],
    [
        ["ClosedPartnerDeals = 1",
         "PartnerDealID IS NOT NULL AND PartnerSalesStage IN ('Support 100%', 'Deploy 100%', "
         "'MS Win 100%', 'Disengaged 0%', 'Loss 0%')",
         "Identifies deals that reached a terminal stage (positive or negative). Used for deal "
         "completion counts."],
        ["PartnerDealsAcceptance = 1",
         "OpportunityID IS NOT NULL AND PartnerDealID IS NOT NULL AND CRM.PartnerName IS NOT NULL "
         "AND AcceptanceStatus = 'Accepted'",
         "Measures partner engagement quality — did the partner formally accept AND link to an "
         "opportunity with a known CRM account?"],
        ["ISVDealsReporting = 1",
         "OpportunityID IS NOT NULL AND CRM.IsCoSellReadyPartnerAccount = 'true'",
         "Gates ISV deal counting. The partner's CRM account must be flagged co-sell ready. A "
         "partner-level attribute, not a deal-level one."],
        ["IsYTDSubSourceDealCount = 'Yes'",
         "CustomerSegment IN ('EPG', 'SMS&P') AND RegistrationStatus IN ('Approved', 'Passed')",
         "Filters YTD counts to only approved deals in Enterprise (EPG) and SMB (SMS&P) segments."],
        ["IsDealAccepted = 'Yes'",
         "IsDealWithPartnerGTM = 1 AND (ReferralMicrosoftSubStatus = 'Accepted' OR "
         "PartnerAcceptanceSubStatus = 'Accepted')",
         "GTM-filtered acceptance. Both sides can trigger acceptance. This is the SAL (Sales Accepted "
         "Lead) gate."],
        ["IsOpportunityWon = 'Yes'",
         "IsDealWithPartnerGTM = 1 AND (BilledRevenueStatus = 'Won' OR (ConsumptionStatus = 'Closed' "
         "AND Reason = 'Completed'))",
         "Two paths to 'Won': traditional billing (Won) or consumption-based closure. Both require "
         "GTM program participation."],
        ["Category Text",
         "IF IsOpportunityWon='Yes' THEN 'Pipeline Wins-SAL-Partner Deals'; ELIF IsDealAccepted='Yes' "
         "THEN 'SAL-Partner Deals'; ELIF IsDealWithPartnerGTM=1 THEN 'Partner Deals'",
         "Funnel categorization for pipeline reporting. Each deal falls into exactly one tier based "
         "on its furthest progression."],
    ],
    col_widths=[1.6, 3.0, 2.3], first_col_bold=True,
)

heading("Velocity & Performance Metrics", level=2)
lead("Measuring speed and effectiveness through the co-sell funnel.")
table(
    ["Metric", "Formula [Unit]", "Meaning"],
    [
        ["PartnerDealApprovalVelocity", "DATEDIFF(ApprovalDate, CreatedDate) [Days]",
         "Days from deal creation to approval. Zero if not yet approved. Measures internal processing "
         "speed; high values signal bottlenecks in deal review."],
        ["PartnerDealCloseVelocity", "DATEDIFF(DueDate, CreatedDate) [Days]",
         "Total lifecycle duration from creation to close/due date. Measures end-to-end cycle time; "
         "used with ApprovalVelocity to see where time is spent (pre- vs post-approval)."],
        ["AverageMonthsForApproval", "DATEDIFF(ApprovalDate, CreatedDate) / 30.0 [Months]",
         "Same as ApprovalVelocity but in months, for trend reporting in monthly aggregation "
         "dashboards."],
        ["VelocityOfInboundDeals", "DATEDIFF(OpportunityCreatedDate, CreatedDate) [Days]",
         "ONLY for Inbound deals. Microsoft's response time — how quickly an MSX Opportunity was "
         "created after the partner's referral. Minimum 1 day (same-day = 1, not 0). Zero for "
         "Outbound."],
        ["IsAverageApproval / IsAverageClose", "'Yes' if date exists, 'No' otherwise [Filter]",
         "Guard flags that filter deals out of average calculations. A deal without an ApprovalDate "
         "shouldn't pull down the average — these flags ensure only deals with valid dates are "
         "included."],
        ["SharingTypeKey", "IF OpportunityKey = 0 THEN 3 ELSE 1 [Enum]",
         "1 = Deal with linked Opportunity (co-sell with visibility); 3 = Deal without Opportunity "
         "(referral only, no MS engagement yet). Drives the MSX Partner Sharing Categorization "
         "analysis."],
    ],
    col_widths=[1.7, 2.4, 2.8], first_col_bold=True,
)
callout("TRANCHE LOGIC:  PartnerRevenueCD > $1M → 'High' tranche.  |  ≤ $1M → 'Low' tranche.  |  "
        "PRCD (capped revenue) = MIN(PartnerRevenueCD, $2M) — prevents outlier deals from skewing "
        "aggregates.", title="Revenue tranches")

heading("Partner Identity Resolution", level=2)
lead("The chain that maps a raw MPN ID to a reportable partner entity.")
para("Partners register with Microsoft using MPN IDs (Microsoft Partner Network). But a single "
     "partner organization may have multiple MPN IDs across subsidiaries. The CoSell pipeline "
     "resolves this into a unified hierarchy:", space_after=4)
table(
    ["Step", "Stage", "Key", "Description"],
    [
        ["1", "MPNId (Partner Center)", "MPNId", "Raw ID from the partner's referral."],
        ["2", "PartnerMaster", "SourceID lookup", "Maps MPN to PartnerOneID / SubKey."],
        ["3", "CRMPartnerAccount", "PartnerNumber",
         "Fallback: CRM account → PipelinePartnerMaster → PartnerOneID."],
        ["4", "PartnerOne", "PartnerOneID", "Unified org-level partner identity."],
        ["5", "PartnerOneSub", "PartnerOneSubKey", "Subsidiary-level (geography-specific)."],
        ["6", "ReportingKey", "PartnerOneReportingKey",
         "Rolled up via Bridge_CascadedPartnerOne + SubsidiaryID."],
    ],
    col_widths=[0.5, 1.9, 1.8, 2.7], first_col_bold=True,
)
para("Resolution priority (from actual code):", bold=True, space_before=6, space_after=2)
numbered("MPNId_PartnerCenter → PartnerMaster.SourceID → PartnerOneID. If the partner's MPN ID is "
         "found in PartnerMaster, use that PartnerOneID directly.", bold_lead="PRIMARY PATH:  ")
numbered("CRMPartnerAccountKey → CRMPartnerAccount.PartnerNumber → PipelinePartnerMaster.SourceID → "
         "PartnerOneID. If MPN lookup fails, traverse through the CRM account's partner number.",
         bold_lead="FALLBACK PATH:  ")
numbered("PartnerMaster.SubsidiaryName → VWDimGeography.SubsidiaryID → PartnerOneSubKey;  OR CRM "
         "path: PipelinePartnerMaster → ReportingPartnerOne.PartnerOneSubID.",
         bold_lead="SUBSIDIARY MAPPING:  ")
numbered("Bridge_CascadedPartnerOne joins on (PartnerOneKey + SubsidiaryID) → "
         "PartnerOneReportingKey. This is the final reportable identity used in Power BI.",
         bold_lead="REPORTING ROLLUP:  ")

heading("Scorecard Recognition Logic", level=2)
lead("The 3rd-of-month boundary rule that determines when a deal gets fiscal credit.")
para("When a partner deal is approved, it must be credited to a specific fiscal month for scorecard "
     "purposes. The recognition month is NOT simply the month of the approval date — a boundary rule "
     "applies:", space_after=4)
code_block(
    "THE 3RD-OF-MONTH RULE\n"
    "IF DAY(ApprovalDate) > 3 OR MONTH(ApprovalDate) = MONTH(RegistrationDate):\n"
    "    ScorecardRecognitionMonth = FORMAT(ApprovalDate, 'MMMM, yyyy')        # Current month\n"
    "ELSE:\n"
    "    ScorecardRecognitionMonth = FORMAT(ApprovalDate - 1 month, 'MMMM, yyyy')  # Prior month"
)
table(
    ["Scenario", "Evaluation", "Recognized", "Note"],
    [
        ["Deal approved March 15", "Day(15) > 3 → TRUE", "March", "Standard case"],
        ["Deal approved March 2, registered March 1",
         "Day(2) ≤ 3 BUT same month as registration", "March", "Same-month exception"],
        ["Deal approved March 2, registered February 15",
         "Day(2) ≤ 3 AND different month", "February", "Pushed back 1 month"],
        ["Deal approved March 1, registered January 10",
         "Day(1) ≤ 3 AND different month", "February", "Pushed back 1 month"],
    ],
    col_widths=[2.2, 2.3, 1.1, 1.3], first_col_bold=True,
)
callout("WHY THIS EXISTS: Prevents gaming of scorecard timing. Deals approved in the first 3 days of "
        "a month (that were registered in a prior month) are likely 'carried over' approvals and "
        "should credit the prior month. This aligns with Microsoft's fiscal period-close procedures.")

heading("Geography Resolution Chain", level=2)
lead("How customer and partner geography are determined with cascading fallbacks. Geography is not a "
     "single field — it is resolved through a cascade of sources.")
para("Customer Geography (CustomerGeographyKey):", bold=True, color=ACCENT, space_after=2)
numbered("From TPAccount (HQ account mapping via opportunity): Opportunity → TPAccountID → "
         "CustomerHQAccountsDim → GPSubsidiaryID → DimCustomerGeography.", bold_lead="PRIMARY:  ")
numbered("Account Number geography (if TPAccountKey = 0): MSXCRMAccountNumber → Hub_vw_Account → "
         "Country → DimSalesGeography → SubsidiaryID.", bold_lead="FALLBACK A:  ")
numbered("Partner-reported subsidiary (if no account number): PartnerDeal.SubsidiaryID → "
         "DimCustomerGeography.", bold_lead="FALLBACK B:  ")
numbered("Private deals (ReferralSubType = 'private - pc') → CustomerGeographyKey forced to 0 "
         "(unknown).", bold_lead="EXCEPTION:  ")
para("Partner Geography (PartnerGeographyKey):", bold=True, color=ACCENT, space_after=2, space_before=4)
bullet("Resolved via PartnerOneSub mapping: PartnerMaster.SubsidiaryName → VWDimGeography.SubsidiaryID;  "
       "OR CRM path → PipelinePartnerMaster → ReportingPartnerOne.PartnerSubsidiaryID.")
bullet("PartnerGeographyKey = SubsidiaryID from DimGeography_Base.")
bullet("PartnerCustomerGeographyKey: cross-reference of customer geography's SubsidiaryID back "
       "through VWDimGeography. Used when partner and customer geography need to be compared.")
bullet("ReportingGeographyKey: derived from GPSubsidiaryID → DimReportingGeography. This is the "
       "geography used for management reporting rollups.")
callout("Multi-level hierarchy: SubsidiaryID → Country → Region → Area. Separate dimensions exist "
        "for Customer Geography, Partner Geography, Reporting Geography, and Capacity Geography to "
        "support different slicing needs.")

heading("IP CoSell, ISV Connect & Marketplace", level=2)
lead("Partner IP attribution, bundled ACR, and marketplace ACV logic.")
para("IP CoSell (Intellectual Property Co-Sell)", bold=True, color=ACCENT, space_after=2)
para("Tracks deals where a partner's IP solution (app, service) is part of the deal.", space_after=2)
bullet("IsIPPartner: from DimIPPartner — flags if the partner has registered IP in the deal.")
bullet("IPCoSellPartnerOneKey: the PartnerOneKey of the IP-providing partner (may differ from the "
       "deal partner).")
bullet("IPPartnerOneReportingKey: reporting rollup for the IP partner, resolved through "
       "FactSolution → DimSolution → PartnerMaster chain.")
bullet("IsPrioritizedPartner: from MapSolutionPracticeIndustryCountry — partner has prioritized "
       "solutions.")
callout("SNAPSHOT OVERRIDE: If PartnerDealSolution_Snapshot has data for a deal, its "
        "PartnerOneReportingKey takes precedence over the solution-based mapping (COALESCE logic).")
para("Marketplace & Bundled Deals", bold=True, color=ACCENT, space_after=2, space_before=4)
bullet("IsMarketplaceRegisteredDeal = 'Yes': PartnerRevenueCD uses MarketPlaceACV from snapshot "
       "(instead of the standard PartnerRevenueCD calculation).")
bullet("IsBundledandPartnerReportedACRDeal = 1 when: Partner reported ACR AND IncentiveType contains "
       "'bundled', OR the deal is marketplace co-sell.")
bullet("D365MonthlyRevenueCD: PartnerRevenueCD / 12 (monthly spread) — ONLY when a D365 EOP/Pipeline "
       "deal + ContractStartDate exists + NOT a marketplace-registered deal.")
bullet("HasD365IncentiveType = 'Yes' when the solution has a D365 incentive ('internal biz apps ip "
       "co-sell', 'msft internal biz apps standard/premium').")
para("ISV Connect", bold=True, color=ACCENT, space_after=2, space_before=4)
para("ISVConnectAppKey links deals to specific ISV Connect apps (DimISVConnectApp). "
     "SellerCoSellIncentiveKey tracks seller incentive programs attached to solutions. ResourceGUID "
     "uniquely identifies Azure resources within IP CoSell deals for ACR tracking. IsIPSameAsCoSell "
     "and IsChannelSameAsCoSell flags identify when the IP/channel partner is the same as the "
     "co-sell partner.")

heading("CoMarketing Module", level=2)
lead("Investment bridge analytics, funnel categorization, and budget tracking. CoMarketing is a "
     "specialized sub-domain tracking joint marketing investments and their ROI across solution "
     "areas, industries, and geographies.")
table(
    ["Area", "Object", "Purpose"],
    [
        ["Budget Allocation", "BudgetProgram",
         "Tracks TPM budget programs for co-marketing investments and partner funding allocations. "
         "Links investment dollars to specific programs."],
        ["Pipeline Forecasting", "ForecastAmount",
         "Forecast amounts for co-sell pipeline and revenue projections across time periods. Used "
         "for budget-vs-actual comparisons."],
        ["Deal Staging", "FunnelCategory / FunnelCategory2",
         "Categorizes deals into funnel stages for pipeline analysis and conversion tracking. Two "
         "levels of categorization for different reporting granularities."],
        ["Partner Funding", "InvestmentAsk",
         "Partner investment requests and approved amounts for co-marketing programs. Tracks the "
         "ask-vs-approved gap."],
        ["Investment Analytics", "InvestBridge (Area / Industry / SolutionArea / SolutionPlay)",
         "Four-dimensional bridge tables that slice co-marketing spend by geographic area, industry "
         "vertical, solution area (Azure/M365/D365/Security), and solution play (granular "
         "solution-level)."],
        ["Trend Analysis", "Opportunity_MonthlySnapshot",
         "Monthly point-in-time snapshots of the opportunity pipeline. Enables month-over-month "
         "trend analysis without losing historical state."],
    ],
    col_widths=[1.4, 2.0, 3.5], first_col_bold=True,
)
callout("CoMarketing has its own fact tables (FactOpportunityCD, FactEngagementMilestoneCD, "
        "FactSolutionCD, FactPartnerDealCD) and dimension copies (DimCustomerCD, DimOpportunityCD, "
        "DimTime) that are optimized for investment analysis.")

heading("Data Architecture: Source to Gold", level=2)
lead("What transforms at each Medallion layer and why it matters.")
para("Source systems unified:", bold=True, space_after=2)
bullet("Partner Center (deals, profiles)")
bullet("MSX / CRM (opportunities, leads)")
bullet("Hub Tables (currency, accounts)")
bullet("Revenue Account (ATU mapping)")
bullet("Partner Master (identity mapping)")
table(
    ["Layer", "Role", "Key Tables", "Purpose"],
    [
        ["Bronze", "Raw ingestion, minimal transformation, schema preservation",
         "Cosell_Bronze_CRP_Solution; Cosell_Bronze_SellInCountryCatRP; PartnerMaster (raw)",
         "Landing zone. Exact copy of source for auditability."],
        ["Silver", "Cleansed & conformed, business keys resolved, intermediate calcs",
         "PartnerDeal_Intermediate; DimOpportunity / _int; Currency / ExchangeRate; "
         "PartnerMaster (mapped); PartnerImpactNumber",
         "Single version of truth for each source entity."],
        ["Gold", "Curated dimensional model, 80+ tables (Dim + Fact), USD-converted financials",
         "FactPartnerDeal (central fact); DimPartnerDeal / DimPartnerOne; DimOpportunity / "
         "DimCustomer; all classification flags computed; all velocity metrics; CoMarketing tables",
         "Power BI-ready star schema. All business logic applied."],
    ],
    col_widths=[0.8, 1.8, 2.7, 1.6], first_col_bold=True,
)
para("Pipeline orchestration & environment management:", bold=True, space_before=6, space_after=2)
bullet("GetNotebookStatus() — checks if a notebook already ran (0 = skip, -1 = error, 2 = "
       "conditional skip). Prevents duplicate processing.")
bullet("GetWorkspaceIDLakehouseID(StreamName) — dynamically resolves workspace + lakehouse IDs per "
       "environment. The same code runs in DEV/UAT/PROD.")
bullet("GetPublishSchema() — returns a versioned schema path. Enables blue-green deployments.")
bullet("writeTable() / getDataframe() — standardized I/O with environment-aware paths.")
bullet("Z-Order optimization applied to reporting tables for query performance. CI/CD via "
       "CoSell_Build_1ES.yml.")

heading("Key Business Questions CoSell Answers", level=2)
lead("The reporting scenarios this data powers.")
para("Partner Performance", bold=True, color=ACCENT, space_after=2)
bullet("How many deals did each partner close this fiscal year?")
bullet("What is the average approval velocity by partner segment?")
bullet("Which partners have the highest deal acceptance rate?")
bullet("What is the deal value distribution across High vs Low tranches?")
para("Revenue & Financial", bold=True, color=ACCENT, space_after=2, space_before=4)
bullet("What is total deal value (USD) by solution area and geography?")
bullet("How does partner-reported revenue compare to Microsoft-verified ACR?")
bullet("What is the marketplace ACV contribution to partner revenue?")
bullet("How do CoMarketing investments correlate with deal pipeline?")
para("Operational Efficiency", bold=True, color=ACCENT, space_after=2, space_before=4)
bullet("How fast is Microsoft responding to inbound referrals?")
bullet("What % of referrals convert to linked MSX opportunities?")
bullet("How many deals are stuck without opportunities (PartnerDealsWithoutOppty)?")
bullet("What is the scorecard recognition distribution by fiscal month?")
para("Strategic & Program", bold=True, color=ACCENT, space_after=2, space_before=4)
bullet("Which ISV co-sell-ready partners are driving the most deals?")
bullet("What is the Partner Impact Number (PIN) across product lines?")
bullet("Are prioritized solution partners outperforming non-prioritized ones?")
bullet("How does the D365 incentive program affect deal closure rates?")

heading("Summary & Key Takeaways", level=2)
takeaways = [
    "CoSell is Microsoft's partner co-selling analytics backbone — it answers 'did the partner deal "
    "work, and how fast?'",
    "Currency conversion uses a date-dependent fallback strategy (Oct 2018 cutoff) with up to 3 "
    "exchange rate sources per metric.",
    "Deal classification has 7 distinct flags, each with specific business rules — they are NOT "
    "interchangeable.",
    "Partner identity resolves through a 6-step chain from raw MPN ID to reportable "
    "PartnerOneReportingKey.",
    "The 3rd-of-month scorecard rule prevents fiscal gaming and aligns with Microsoft's period-close "
    "procedures.",
    "Geography cascades through 3 fallback sources with special handling for private (PC) deals.",
    "IP CoSell and Marketplace deals have separate revenue attribution logic (snapshot overrides, "
    "ACV substitution, D365 monthly spread).",
    "The pipeline is fully environment-portable — the same code runs in DEV/UAT/PROD via dynamic "
    "workspace resolution.",
]
for i, tk in enumerate(takeaways, start=1):
    numbered(tk)

# =========================================================================
# PART 3 — CRM COSELL DEEP DIVE
# =========================================================================
heading("Part 3 — CRM Co-sell Deep Dive: Data & Technical", level=1, page_break=True)
lead("Source deck: GPS Domain Training — CRM Cosell Deep Dive. The underlying data sources, GPS Mart "
     "views and tables, the star-schema ER model, important analytical queries, the CRM Tabular "
     "model, and Marketplace security implementation.")

heading("Upstream Sources & GPS Mart", level=2)
para("The GPS Mart is the SQL Endpoint for users to pull data for ad-hoc reporting asks and "
     "analysis.", space_after=2)
table(
    ["Property", "Value"],
    [
        ["GPS Mart — Server", "gpsmart.database.windows.net"],
        ["GPS Mart — ADW / Database", "GPSMart"],
        ["ACR-related data", "Found in the azure schema in GPS Mart"],
    ],
    col_widths=[2.2, 4.7], first_col_bold=True,
)
para("Key azure-schema views in GPS Mart:", bold=True, space_before=4, space_after=2)
bullet("vw_Bridge_AzureConsumptionPartner — holds the ACR partner association data. The underlying "
       "table is Map_Association_Partner_PPR_OCPMart (a different table than the Map Association "
       "Partner table projects use from ADLS; it has limited columns and PPR-specific filters).",
       bold_lead="")
bullet("vw_Fact_AzureConsumption — revenue/consumption data along with the keys that are additive "
       "for ACR.")
bullet("Dimension views: vw_Dim_AZPricingLevel (Azure-sales-specific pricing levels), "
       "vw_Dim_AzureFRAGroup (Azure FRA Group data), plus Azure Offer, Azure Pricing Level, Azure "
       "Service Level, Azure Subscription, and Azure Service Influencer.")

heading("GPS Mart Views and Tables", level=2)
para("GPS Mart Connection — Server: gpsmart.database.windows.net, ADW: GPSMart.", italic=True,
     size=10, color=GREY, space_after=4)
table(
    ["View Name (Schema: crm)", "Table Name", "Primary Key"],
    [
        ["vw_Dim_Opportunity", "DimOpportunity", "OpportunityKey"],
        ["vw_Fact_Opportunity", "FactOpportunity", "Composite Primary Key"],
        ["vw_Dim_PartnerDeal", "DimPartnerDeal", "PartnerDealKey"],
        ["vw_Fact_PartnerDeal", "FactPartnerDeal", "Composite Primary Key"],
        ["vw_Dim_Solution", "DimSolution", "SolutionKey"],
        ["vw_Fact_Solution", "FactSolution", "Composite Primary Key"],
        ["vw_Dim_EngagementMilestone", "DimEngagementMilestone", "EngagementMilestoneKey"],
        ["vw_Fact_EngagementMilestone", "FactEngagementMilestone", "Composite Primary Key"],
        ["vw_Dim_MarketplaceInvoice", "DimMarketplaceInvoice", "InvoiceKey"],
        ["vw_Fact_MarketplaceBilledSales", "FactMarketplaceBilledSales", "Composite Primary Key"],
    ],
    col_widths=[2.7, 2.6, 1.6],
)

heading("GPS Mart ER Diagram (Star Schema)", level=2)
para("This is a standard star schema: the fact tables are surrounded by and related to the "
     "dimension tables around them, except the bridge table which is further related for partner "
     "data. Every key in a fact table has a related dimension; those keys are the join columns.",
     space_after=4)
para("Core relationships (fact → keys):", bold=True, space_after=2)
table(
    ["Fact / View", "Joins On Keys"],
    [
        ["vw_Fact_Opportunity", "OpportunityKey, EngagementMilestoneKey"],
        ["vw_Fact_PartnerDeal", "PartnerDealKey, OpportunityKey, SolutionKey"],
        ["vw_Fact_EngagementMilestone", "EngagementMilestoneKey"],
        ["vw_Fact_MarketplaceBilledSale", "OfferKey, PartnerDealKey"],
    ],
    col_widths=[2.6, 4.3], first_col_bold=True,
)
para("Related dimensions: vw_Dim_Solution (SolutionKey), vw_Dim_Opportunity (OpportunityKey), "
     "vw_Dim_EngagementMilestone (EngagementMilestoneKey), vw_Dim_PartnerDeal (PartnerDealKey), "
     "vw_Dim_MarketplaceOffer (OfferKey).", space_after=4)
para("Joining examples (from the training):", bold=True, space_after=2)
bullet("To see consumption by AZ Pricing Level 2: join the dimension view and the fact view on the "
       "AZ Pricing Level 2 key.")
bullet("To see data for Azure FRA Group 2: join vw_Dim_AzureFRAGroup and the fact view on the Azure "
       "FRA Group key (column Dim_AzureFRAGroupID). The same applies to every dimension view related "
       "directly to the fact view.")
bullet("vw_Bridge_AzureConsumptionPartner is a bridge view used to join partner-related dimensions "
       "(vw_Dim_PartnerAttachType and vw_ReportingPartnerOne). To see data at partner level, join "
       "the fact table to the bridge table, and the bridge table to either DimPartnerAttachType or "
       "ReportingPartnerOne.")
callout("Additivity rule: attributes related DIRECTLY to the fact table are always additive; "
        "attributes that need a BRIDGE table are never additive. That is precisely why bridge tables "
        "are added to the data model — you do not need to memorize additivity, just follow the join "
        "path.", title="Additive vs non-additive")

heading("Important Queries", level=2)
para("Marketplace Billed Sales (MBS) — Co-sell and Total:", bold=True, space_after=2)
code_block(
    "-- MBS Co-sell\n"
    "WITH CTE_MBS AS (\n"
    "  SELECT DocumentID, ProductID, ServicePeriodDuration,\n"
    "         PartnerOneReportingID, TPAccountID,\n"
    "         MAX(TotalPriceCDAmountAnnualizedwithCap) AS MBSCosell\n"
    "  FROM crm.vw_Fact_MarketplaceBilledSales\n"
    "  GROUP BY DocumentID, ProductID, ServicePeriodDuration,\n"
    "           PartnerOneReportingID, TPAccountID\n"
    ")\n"
    "SELECT SUM(MBSCosell) FROM CTE_MBS;",
    label="MBS Co-sell"
)
code_block(
    "-- Total MBS\n"
    "WITH CTE AS (\n"
    "  SELECT DISTINCT InvoiceKey, TotalPriceCDAmount\n"
    "  FROM crm.vw_Fact_MarketplaceBilledSales\n"
    ")\n"
    "SELECT SUM(TotalPriceCDAmount) AS TotalMBS FROM CTE;",
    label="Total Marketplace Billed Sales"
)
para("ACR — Partner Reported (PRACR) and Deal Reg (DRACR):", bold=True, space_before=4, space_after=2)
code_block(
    "-- PRACR (Partner Reported ACR)\n"
    "WITH CTE AS (\n"
    "  SELECT DISTINCT FPD.PartnerDealKey, FPD.ResourceGUID,\n"
    "         FPD.ApprovalAndReportedDate, FPD.TrueACRConsumption, DPD.PSXDealID\n"
    "  FROM [crm].[vw_Fact_PartnerDeal] FPD\n"
    "  INNER JOIN [crm].[vw_Dim_PartnerDeal] DPD ON DPD.PartnerDealKey = FPD.PartnerDealKey\n"
    "  WHERE [IsPartnerReportedACR] = 'Yes'\n"
    "    AND FPD.[ApprovalAndReportedFiscalYear] = 'FY24'\n"
    ")\n"
    "SELECT SUM(TrueACRConsumption) AS PRACR FROM CTE;",
    label="Partner Reported ACR"
)
code_block(
    "-- DRACR (Deal Reg ACR)\n"
    "WITH CTE AS (\n"
    "  SELECT DISTINCT FPD.PartnerDealKey, FPD.[AzureIPCosellDealReg]\n"
    "  FROM [crm].[vw_Fact_PartnerDeal] FPD\n"
    "  INNER JOIN [crm].[vw_Dim_PartnerDeal] DPD ON DPD.PartnerDealKey = FPD.PartnerDealKey\n"
    "  WHERE DPD.[IsAzurePartnerCosellDealEOP] = 'Yes'\n"
    ")\n"
    "SELECT SUM([AzureIPCosellDealReg]) AS DRACR FROM CTE;",
    label="Deal Reg ACR"
)

heading("CRM Tabular Model", level=2)
table(
    ["Property", "Value"],
    [
        ["CRM Tabular — Server", "GPSInsightsASCRM"],
        ["CRM Tabular — Database", "CRM Tabular"],
    ],
    col_widths=[2.2, 4.7], first_col_bold=True,
)
para("Important measure groups / dimensions:", bold=True, space_before=4, space_after=2)
bullet("Partner Deal")
bullet("Opportunity")
bullet("Solution")
bullet("Marketplace Billed Sales")
bullet("Customer")
bullet("Partner Sharing")

heading("Marketplace Security Implementation", level=2)
lead("MBS row-level security example — how different user roles see different slices of the same "
     "data, based on Customer-flagged (TPID) and MS Sales (Subsidiary) access.")
para("Sample data set:", bold=True, space_after=2)
table(
    ["PartnerOne", "TPID", "Subsidiary", "MBS"],
    [
        ["P1", "T2", "S1", "$1,000"],
        ["P1", "T1", "S2", "$2,000"],
        ["P2", "T2", "S5", "$3,000"],
        ["P3", "T1", "S4", "$4,000"],
        ["P4", "T2", "S1", "$5,000"],
        ["P5", "T2", "S3", "$6,000"],
        ["P5", "T3", "S2", "$7,000"],
        ["P5", "T5", "S4", "$8,000"],
        ["P6", "T4", "S5", "$9,000"],
    ],
    col_widths=[1.6, 1.6, 1.6, 1.6],
)
para("Access role assignments:", bold=True, space_before=4, space_after=2)
table(
    ["Users with Customer-flagged role (TPID)", "Users with MS Sales access (Subsidiary)"],
    [
        ["U2 → T1", "U3 → S1"],
        ["U4 → T2", "U4 → S2"],
    ],
    col_widths=[3.4, 3.5],
)
para("Resulting visibility per user:", bold=True, space_before=4, space_after=2)
table(
    ["U2 sees (TPID = T1)", "", "", ""],
    [
        ["PartnerOne", "TPID", "Subsidiary", "MBS"],
        ["P1", "T1", "S2", "$2,000"],
        ["P3", "T1", "S4", "$4,000"],
    ],
    col_widths=[1.6, 1.4, 1.7, 1.6],
)
table(
    ["U3 sees (Subsidiary = S1)", "", "", ""],
    [
        ["PartnerOne", "TPID", "Subsidiary", "MBS"],
        ["P1", "T2", "S1", "$1,000"],
        ["P4", "T2", "S1", "$5,000"],
    ],
    col_widths=[1.6, 1.4, 1.7, 1.6],
)
table(
    ["U4 sees (TPID = T2 OR Subsidiary = S2)", "", "", ""],
    [
        ["PartnerOne", "TPID", "Subsidiary", "MBS"],
        ["P1", "T2", "S1", "$1,000"],
        ["P2", "T2", "S5", "$3,000"],
        ["P4", "T2", "S1", "$5,000"],
        ["P5", "T2", "S3", "$6,000"],
    ],
    col_widths=[1.6, 1.4, 1.7, 1.6],
)

# =========================================================================
# PART 4 — GPS 301 PARTNER SHARING
# =========================================================================
heading("Part 4 — GPS 301: Partner Sharing Report", level=1, page_break=True)
lead("Source deck: GPS 301 — Partner Sharing (Chetan). The Partner Sharing report, its Billed and "
     "Consumption pipeline metrics, the FY26 metric definitions, predefined slicers, refresh "
     "architecture, and the reporting glossary.")

heading("Partner Sharing Report Overview", level=2)
para("Agenda of the Partner Sharing session:", bold=True, space_after=2)
bullet("Partner Sharing Report")
bullet("Architecture")
bullet("Refresh Frequency")
bullet("Optimization Techniques")
para("The report covers the Partner Sharing pipeline across both Billed and Consumption "
     "opportunities, with predefined slicers that enforce the Partner Sharing metric logic.",
     space_after=4)

heading("Billed Pipeline Metrics", level=2)
term("# Billed Oppty", "The total number of billed opportunities, calculated by filtering for "
     "records classified as billed (TypeID = 1), while preserving all existing report filters to "
     "maintain contextual accuracy.")
term("Partner Sharing % (Billed Oppty)", "Percentage of billed opportunities that include partner "
     "sharing, measured against total billed opportunities.")
term("Co-Sell Prioritized Sharing % (Billed Oppty)", "Percentage of billed co-sell prioritized "
     "opportunities that involve partner sharing, relative to total billed co-sell prioritized "
     "opportunities.")

heading("Consumption Pipeline Metrics", level=2)
term("# Consumption Oppty", "Total value of consumption opportunities, aggregated across all "
     "opportunities classified as consumption.")
term("Partner Sharing % (Consumption Oppty)", "Percentage of consumption opportunities that include "
     "partner sharing, measured against total consumption opportunities.")
term("Co-Sell Prioritized Sharing % (Consumption Oppty)", "Percentage of consumption opportunities "
     "that are completed, co-sell prioritized, and include partner sharing, measured against total "
     "completed co-sell prioritized consumption opportunities.")

heading("FY26 Partner Sharing Metrics", level=2)
lead("Latest metric definitions (Partner Sharing — Web view).")
table(
    ["Metric", "Definition"],
    [
        ["Partner Sharing % ($)",
         "Measures the proportion of opportunity value that is shared with partners, across both "
         "Billed and Consumption opportunities."],
        ["C2C Shared %",
         "Measures how much partner-shared, committed consumption has actually been completed, based "
         "on milestone value."],
        ["U2C Shared %",
         "Measures the proportion of committed, partner-shared consumption opportunity value relative "
         "to qualified in-progress consumption opportunities."],
        ["Shared Won/Completed %",
         "Measures the proportion of won (billed) and completed (consumption) opportunities that are "
         "partner-engaged, where Partner-engaged = Shared + Associated."],
    ],
    col_widths=[2.0, 4.9], first_col_bold=True,
)

heading("Reporting Structure, Slicers & Architecture", level=2)
para("Summary by Solution Area", bold=True, color=ACCENT, space_after=2)
para("Opportunities are reported across different solution areas using a hierarchical structure, "
     "which can be expanded to view detailed solution area breakdowns.", space_after=4)
para("Predefined Slicers", bold=True, color=ACCENT, space_after=2)
para("As per the Partner Sharing requirements, all necessary slicers have already been applied to "
     "the report to ensure consistency with the defined metric logic. These slicers filter the data "
     "based on key conditions such as qualified pipeline, business type, opportunity status, and "
     "account tagging. This ensures the insights presented are aligned with the Partner Sharing "
     "framework and reflect only the relevant subset of opportunities for accurate analysis and "
     "reporting.", space_after=4)
para("Architecture & Refresh", bold=True, color=ACCENT, space_after=2)
para("The deck covers the Partner Sharing reporting architecture and the Partner Sharing Reports "
     "refresh timeline, built on the Medallion architecture and reporting lakehouse described in the "
     "glossary below.", space_after=4)

heading("Reporting Architecture Glossary", level=2)
table(
    ["Term", "Definition"],
    [
        ["Medallion Architecture",
         "A layered data architecture that processes data through multiple stages — typically bronze, "
         "silver, and gold layers — to ensure data quality and readiness for analysis."],
        ["Reporting Lakehouse",
         "A unified data platform that combines the features of data lakes and data warehouses to "
         "store and manage data used for reporting and analytics."],
        ["Semantic Models",
         "Abstracted representations of data that simplify complex data structures, making it easier "
         "for users to understand and analyze data efficiently."],
        ["Data Processing",
         "The series of operations performed on raw data to convert it into a usable format for "
         "analysis and reporting."],
    ],
    col_widths=[1.8, 5.1], first_col_bold=True,
)

# Closing
divider()
para("End of compendium — consolidated from GPSCRM Co-sell, CoSell Domain Business Logic Guide, "
     "CRM Cosell Deep Dive, and GPS 301 Partner Sharing. Global Partner Solutions | MCAPS Data "
     "Engineering | FY26.", italic=True, size=10, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

# ---- Save ----------------------------------------------------------------
os.makedirs(os.path.dirname(OUT_DOCX_COPY), exist_ok=True)
doc.save(OUT_DOCX_COPY)
try:
    doc.save(OUT_DOCX)
    print(f"Saved primary: {OUT_DOCX}")
except Exception as e:
    print(f"WARN could not save to Desktop path: {e}")
print(f"Saved workspace copy: {OUT_DOCX_COPY}")
print(f"Total paragraphs: {len(doc.paragraphs)} | Tables: {len(doc.tables)}")
