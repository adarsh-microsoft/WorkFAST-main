"""
Build the FY27 IP Co-Sell Transition Approach Document from the Co-Marketing template.
- Preserves page 1 (cover/info) format exactly (only updates the PROJECT INFO ADO Item value).
- Replaces the manual Table of Contents with a live, auto-updating TOC field.
- Rebuilds the body (page 3+) with FY27 IP Co-Sell content using the template's styles.
"""
import copy
import re
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = r"c:\WorkFAST-main\generated-content\approach-documents\Approach Document - Co-Marketing v1.2.docx"
OUT = r"c:\WorkFAST-main\generated-content\approach-documents\Approach Document - FY27 IP Co-Sell Transition v1.0.docx"

doc = Document(SRC)
body = doc.element.body

# ---------------------------------------------------------------------------
# 1) PAGE 1: update PROJECT INFO -> ADO Item value (format preserved)
# ---------------------------------------------------------------------------
NEW_ADO = "Business Scenario 49754: FY27 IP Co-Sell Transition \u2014 Compensation Framework & DRACR Updates"
updated_ado = False
for tbl in doc.tables:
    head = tbl.rows[0].cells[0].text.strip().upper()
    if head.startswith("PROJECT INFO"):
        for row in tbl.rows:
            if row.cells[0].text.strip().upper().startswith("ADO ITEM"):
                cell = row.cells[1]
                para = cell.paragraphs[0]
                if para.runs:
                    # keep first run's formatting, drop the rest
                    para.runs[0].text = NEW_ADO
                    for r in list(para.runs)[1:]:
                        r._element.getparent().remove(r._element)
                else:
                    para.add_run(NEW_ADO)
                updated_ado = True
print("PROJECT INFO ADO Item updated:", updated_ado)

# ---------------------------------------------------------------------------
# 2) Locate "Table of Contents" heading + the manual TOC paragraph after it
# ---------------------------------------------------------------------------
from docx.text.paragraph import Paragraph

paras = doc.paragraphs
toc_heading_idx = None
for i, p in enumerate(paras):
    if (p.style and p.style.name == "Heading 2") and p.text.strip().lower() == "table of contents":
        toc_heading_idx = i
        break
assert toc_heading_idx is not None, "Could not find 'Table of Contents' heading"

# Exclude the "Table of Contents" heading itself from the TOC field while keeping a
# heading-like appearance. Word strips a *direct* outlineLvl override on save, so use a
# dedicated 'TOC Heading' paragraph style (based on Heading 2) with a style-level
# outlineLvl=9 (body text) — which Word preserves and excludes from the TOC.
if 'TOC Heading' not in [s.name for s in doc.styles]:
    _ts = doc.styles.add_style('TOC Heading', WD_STYLE_TYPE.PARAGRAPH)
    _ts.base_style = doc.styles['Heading 2']
    _sel = _ts.element
    _spPr = _sel.find(qn('w:pPr'))
    if _spPr is None:
        _spPr = OxmlElement('w:pPr')
        _sel.insert_element_before(_spPr, 'w:rPr', 'w:tblPr', 'w:trPr', 'w:tcPr', 'w:tblStylePr')
    _sol = OxmlElement('w:outlineLvl')
    _sol.set(qn('w:val'), '9')
    _spPr.insert_element_before(_sol, 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr', 'w:pPrChange')
paras[toc_heading_idx].style = doc.styles['TOC Heading']

# manual TOC = next non-empty paragraph after the heading
toc_para = None
for p in paras[toc_heading_idx + 1:]:
    if p.text.strip():
        toc_para = p
        break
assert toc_para is not None, "Could not find manual TOC paragraph"
print("Manual TOC text (first 60):", repr(toc_para.text[:60]))


def make_toc_field(paragraph):
    """Convert a paragraph into a live TOC field (levels 1-3)."""
    p = paragraph._p
    for r in list(paragraph.runs):
        r._element.getparent().remove(r._element)

    def add_run(child):
        r = OxmlElement('w:r')
        r.append(child)
        p.append(r)

    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    begin.set(qn('w:dirty'), 'true')
    add_run(begin)

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    add_run(instr)

    sep = OxmlElement('w:fldChar')
    sep.set(qn('w:fldCharType'), 'separate')
    add_run(sep)

    t = OxmlElement('w:t')
    t.text = 'Right-click here and choose "Update Field" to build the Table of Contents.'
    add_run(t)

    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    add_run(end)


make_toc_field(toc_para)
toc_p_el = toc_para._p

# ---------------------------------------------------------------------------
# 3) Remove every body block AFTER the TOC field (keep trailing sectPr)
# ---------------------------------------------------------------------------
removing = False
for child in list(body):
    if child is toc_p_el:
        removing = True
        continue
    if removing:
        if child.tag == qn('w:sectPr'):
            continue
        body.remove(child)

# ---------------------------------------------------------------------------
# 4) Helpers for appending styled content (appends before trailing sectPr)
# ---------------------------------------------------------------------------

def H1(text, page_break=False):
    p = doc.add_paragraph(text, style="Heading 1")
    if page_break:
        p.paragraph_format.page_break_before = True
    return p

def H2(text):
    return doc.add_paragraph(text, style="Heading 2")

def H3(text):
    # strip trailing milestone tags like " [M3, M5]" / " [M6]"
    text = re.sub(r'\s*\[M\d+(?:,\s*M\d+)*\]\s*$', '', text)
    return doc.add_paragraph(text, style="Heading 3")

def P(text):
    return doc.add_paragraph(text, style="Normal")

def B(text):
    return doc.add_paragraph(text, style="List Bullet")

def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(h)
        run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return t

# ---------------------------------------------------------------------------
# 5) BODY CONTENT (page 3+)
# ---------------------------------------------------------------------------
H1("Approach Document", page_break=True)

H2("BRD Location")
P("FY27 IP Co-Sell Requirements v1.0 (2026-04-01) \u2014 Aligned across business "
  "(Compliance, Operations, Finance, Marketplace, WWIC, GPO).")

H2("Business Request -")
P("For FY27, unify Marketplace and non-Marketplace IP Co-Sell deal recognition into a single seller "
  "compensation framework. Marketplace IP Co-Sell deals remain eligible for credit via MBS; "
  "non-Marketplace deals remain eligible for seller credit via DRACR, capped per deal for risk "
  "containment. Eligibility is restricted to a curated Top-25 IP partner list (locked for FY27) plus "
  "SAP as a PRACR exception. PRACR is broadly retired, with SAP continuing under the renamed "
  "\u201cSAP Tenant Consumption.\u201d Both credit types land in the total ACR bucket for quota "
  "retirement. FY27 Finance targets, Service Comp Group (SCG) taxonomy, and fiscal date boundaries "
  "are refreshed. Hard launch deadline: July 1, 2026 (FY27 Q1 start).")

H2("Approach -")
P("The work is decomposed into two functional workstreams \u2014 Data Model & Pipeline (7 requirements) "
  "and Reporting (9 requirements) \u2014 executed under the standard SDLC track. Data engineering changes "
  "are parameter- and join-level only on the existing FactIPCoSell pipeline (no architectural rewrite), "
  "and reporting changes update existing surfaces only (no new visuals) to manage launch cognitive load.")

# ---- Data Model & Pipeline -------------------------------------------------
H2("Data Model & Pipeline - Functional Approach")

H3("1. MBS vs DRACR Credit-Path Routing [M3, M5]")
B("Classify every IP Co-Sell deal at deal grain by Marketplace vs non-Marketplace.")
B("Route Marketplace deals to MBS credit; route non-Marketplace deals to DRACR (IP Co-Sell) credit.")
B("Land both credit paths in the unified total ACR bucket for quota retirement.")
B("Apply the Top-25 partner eligibility filter (plus SAP exception) at the routing layer; credit DRACR "
  "only when the IP partner matches an eligible/exception partner.")

H3("2. SAP Exception Path - Rename PRACR to \u201cSAP Tenant Consumption\u201d [M5, M6]")
B("Rename the existing PRACR pipeline, fields, and measures to \u201cSAP Tenant Consumption.\u201d")
B("Flag SAP reporting as offline/exception and isolate it from the standard Top-25 flow.")
B("Handle dual-credit de-duplication so SAP volume is not double-counted across DRACR + MBS.")
B("Callout: SAP dual-credit continuation is pending formal metric definition and compliance sign-off; "
  "additional effort may be required if scope changes.")

H3("3. FY27 IP Co-Sell Target Integration [M5, M6]")
B("Integrate FY27 IP Co-Sell target data from Finance (Ben Frisbee / Nathan Taylor) into dashboard measures.")
B("Map targets to MetricKey=3 (IP Co-Sell Azure) FY27 planning numbers.")
B("Callout: FY27 target source system and grain must be confirmed with stakeholders before dashboard integration.")

H3("4. New Rejection Reason Codes in ReasonsCalculatedTable [M5]")
B("Add reason codes: \u201cPartner not in Top 25,\u201d \u201cACV below $25K,\u201d "
  "and \u201cPRACR ineligible (non-SAP).\u201d")
B("Surface reason codes for excluded/rejected deals to support seller transparency and operations triage.")

H3("5. Duplicate Detection in FactIPCoSell [M5]")
B("Implement deal-grain duplicate detection to prevent a single deal from receiving DRACR + MBS + "
  "SAP Tenant Consumption credit simultaneously.")
B("Define and apply a credit-precedence rule for deals qualifying under more than one path "
  "(see Clarification).")

H3("6. FY27 Program Changes [M5]")
B("Update FactIPCoSell parameters: Bundled 30% \u2192 50% / $300K \u2192 $500K cap; BYOL retained at "
  "5% / $50K; add $15K floor for ACV \u2265 $25K.")
B("Expand partner ingestion from 10 \u2192 25 MarketplaceTransitionPartners (plus SAP exception); "
  "partner list locked for FY27 \u2014 no mid-year additions.")
B("Replace hardcoded FY26 dates (e.g., '2025-07-01') with FY27 fiscal filters and date boundaries "
  "across affected notebooks.")

H3("7. SCG Remap & FY27 Exchange Rates [M5]")
B("Remap Service Comp Groups (SCG) to the FY27 taxonomy (mapping sourced from the FD&E team \u2014 "
  "Padmaja's team), backward-compatible for FY26 closeout reporting.")
B("Load FY27 dollar exchange rates for currency conversion in credit calculations.")

# ---- Reporting -------------------------------------------------------------
H2("Reporting - Functional Approach")

H3("1. Rename \u201cTransition Partner Performance\u201d Tab to \u201cIPCS Partner Performance\u201d [M6]")
B("Rename the existing \u201cTransition Partner Performance\u201d tab to \u201cIPCS Partner Performance.\u201d")

H3("2. Dedicated SAP Tab [M6]")
B("Separate SAP from the 25 IP Co-Sell partners into a dedicated SAP tab.")

H3("3. Retain PRACR Metrics & Deal Registration on the SAP Tab [M6]")
B("Retain PRACR metrics and deal-registration visuals on the SAP tab (under SAP Tenant Consumption).")

H3("4. Mirror IPCS Structure with SAP-Specific Filters [M6]")
B("Mirror the IPCS Partner Performance structure on the SAP tab with SAP-specific filters.")

H3("5. Trim PRACR Metrics from the 25-Partner (IPCS) Tab [M6]")
B("Remove PRACR-related metrics and visuals from the 25-partner (IPCS) tab; retain deal registration "
  "and target metrics only.")

H3("6. Remove \u201cDeal Registration by Incentive Type\u201d Visual [M6]")
B("Remove the \u201cDeal Registration by Incentive Type\u201d visual.")

H3("7. Retain \u201cDeal Registration by Segment\u201d and \u201cDeal Direction\u201d Visuals [M6]")
B("Retain the \u201cDeal Registration by Segment\u201d and \u201cDeal Direction\u201d visuals on the IPCS tab.")

H3("8. Remove Biz Apps Performance & Pipeline Tabs [M6]")
B("Remove the Biz Apps Performance and Pipeline tabs.")
B("Callout: Biz Apps tab removal to be communicated to the Biz Apps team by Savvy; escalations directed "
  "to Antoine Boris.")

H3("9. Integrate FY27 IP Co-Sell Target Data into Dashboard Measures [M6]")
B("Integrate FY27 IP Co-Sell target data from Finance (Ben Frisbee / Nathan Taylor) into dashboard measures.")

# ---- Assumptions & Callouts ------------------------------------------------
H2("Assumptions & Callouts")
B("DRACR core logic and the 26-partner implementation are already in place; no changes expected post-May.")
B("PRACR is retired for all partners except SAP; SAP continues under the renamed \u201cSAP Tenant Consumption.\u201d")
B("SAP dual-credit continuation is pending formal metric and compliance sign-off; if confirmed, additional "
  "effort may be required depending on scope change.")
B("Partner list is locked for FY27 \u2014 no mid-year additions.")
B("ACV data source (Partner Revenue) remains the same \u2014 sourced from ReferralMasterPII.")
B("DCF S10 scoping is purely operational and requires no change in logic.")
B("FY27 target data source and grain must be confirmed with stakeholders before dashboard integration.")
B("Biz Apps tab removal to be communicated to the Biz Apps team by Savvy; escalations directed to Antoine Boris.")
B("DRACR should be credited only if the IP partner matches one of the exception partners.")
B("SCG mapping will be required from the FD&E team (Padmaja's team).")

# ---- Source Mapping --------------------------------------------------------
H2("Source Mapping")
add_table(
    ["Source", "Target", "Field / Measure"],
    [
        ["Partner Center deal registration", "FactIPCoSell", "Marketplace vs non-Marketplace deal flag (MBS vs DRACR routing)"],
        ["MBS deal feed", "Unified ACR bucket", "IP Co-Sell credit (Marketplace path)"],
        ["FactIPCoSell (DRACR calc)", "Unified ACR bucket", "IP Co-Sell credit (non-Marketplace, capped)"],
        ["PRACR pipeline", "SAP Tenant Consumption pipeline", "SAP exception credit (renamed); offline/exception flag"],
        ["ReferralMasterPII", "FactIPCoSell", "Partner Revenue (ACV) \u2014 unchanged"],
        ["EAC vetting status feed", "Top-25 eligibility join", "IsEligiblePartner flag + audit trail"],
        ["MarketplaceTransitionPartners Excel (25)", "Partner ingestion", "Top-25 partner list (locked FY27)"],
        ["FY27 Target data (Finance: Frisbee / Taylor)", "Semantic model measures", "MetricKey=3 (IP Co-Sell Azure) FY27 targets"],
        ["SCG mapping (FD&E / Padmaja's team)", "SCG remap step", "FY27 Service Comp Group taxonomy"],
        ["FY27 $ exchange rates", "Pipeline ingestion", "Currency conversion"],
        ["Rejection logic", "ReasonsCalculatedTable", "\u201cPartner not in Top 25,\u201d \u201cACV below $25K,\u201d \u201cPRACR ineligible (non-SAP)\u201d"],
    ],
)

# ---- Changes required for Reporting and Downstream -------------------------
H2("Changes required for Reporting and Downstream")
B("Report \u2014 rename Transition Partner Performance \u2192 IPCS Partner Performance; add a dedicated "
  "SAP tab (PRACR metrics + deal registration retained); remove PRACR metrics and the "
  "\u201cDeal Registration by Incentive Type\u201d visual from the IPCS tab; remove Biz Apps Performance "
  "and Pipeline tabs; retain Deal Registration by Segment and Deal Direction visuals.")
B("Semantic model \u2014 unified ACR bucket (MBS + DRACR); SAP Tenant Consumption credit view "
  "(dual-credit pending compliance); per-deal cap visualization; FY27 DAX parameters "
  "(Bundled 50%/$500K, BYOL 5%/$50K, $15K floor); MetricKey=3 FY27 targets; new rejection reason "
  "codes; duplicate-detection measure.")
B("Data \u2014 FactIPCoSell parameter updates (caps/rates/floor, partner list 10 \u2192 25, FY27 date "
  "boundaries); SAP Tenant Consumption pipeline; EAC vetting eligibility join with audit trail; "
  "SCG remap to FY27 taxonomy; FY27 exchange-rate load; duplicate detection across DRACR/MBS/SAP.")
B("Documentation \u2014 data dictionary, IC guide, manager toolkit, compliance runbook, and pipeline "
  "runbook; onboarding and comms cascade.")
B("Acceptance \u2014 written UAT sign-off from WWIC, Finance, Compliance, Ops, and Reporting; zero "
  "open UAT defects at handover; hard launch July 1, 2026.")

# ---- Clarification ---------------------------------------------------------
H2("Clarification-")
B("FY27 IP Co-Sell target data: confirm source system AND grain with Finance (Ben Frisbee / Nathan "
  "Taylor) before dashboard integration.")
B("SAP dual-credit continuation: confirm formal metric definition and compliance sign-off; scope change "
  "may require additional effort.")
B("SCG mapping: confirm delivery and timing of the FY27 SCG taxonomy from the FD&E team (Padmaja's team).")
B("Duplicate-credit de-duplication: confirm the precedence rule when a deal qualifies for more than one "
  "of DRACR / MBS / SAP Tenant Consumption.")
B("DRACR exception-partner match: confirm the authoritative exception-partner list used to gate DRACR crediting.")
B("Biz Apps tab removal: confirm Savvy's communication to the Biz Apps team; escalations to Antoine Boris.")
B("FY27 $ exchange rates: confirm the source and effective date for the FY27 rate load.")

# ---- TestCases -------------------------------------------------------------
H2("TestCases")
P("Planned validation scenarios:")
B("MBS vs DRACR routing correctness at deal grain (Marketplace \u2192 MBS, non-Marketplace \u2192 DRACR).")
B("Per-deal cap and $15K floor enforcement (Bundled 50%/$500K; BYOL 5%/$50K) for ACV \u2265 $25K.")
B("Top-25 partner eligibility + SAP exception gating; DRACR credited only for matched exception partners.")
B("Duplicate-credit prevention \u2014 no deal receives DRACR + MBS + SAP Tenant Consumption simultaneously.")
B("Rejection reason codes populate correctly (\u201cPartner not in Top 25,\u201d \u201cACV below $25K,\u201d "
  "\u201cPRACR ineligible (non-SAP)\u201d).")
B("SAP tab parity \u2014 PRACR metrics + deal registration render on the SAP tab; IPCS tab no longer "
  "shows PRACR metrics or the Incentive Type visual.")
B("FY27 target integration \u2014 MetricKey=3 measures reflect FY27 Finance numbers at the agreed grain.")
B("Sign-off required from WWIC, Finance, Compliance, Ops, and Reporting prior to the July 1, 2026 launch.")

# ---------------------------------------------------------------------------
# 6) Tell Word to update fields (TOC) on open
# ---------------------------------------------------------------------------
settings = doc.settings.element
upd = settings.find(qn('w:updateFields'))
if upd is None:
    upd = OxmlElement('w:updateFields')
    settings.append(upd)
upd.set(qn('w:val'), 'true')

doc.save(OUT)
print("Saved:", OUT)
